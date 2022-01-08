from personnel_discharge import *
import docx
import pandas as pd
import personnel_window
from backup_bd import *
import log_window
import classes
from PyQt5.QtWidgets import QMainWindow
from PyQt5.QtWidgets import QTableWidgetItem


#ОКНО_АДМИНА
class disc_panel(QMainWindow):
    def __init__(self, parent=None):
        super(disc_panel, self).__init__(parent)
        self.ui = pers_discharge()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(lambda: self.write_table())
        self.ui.pushButton_5.clicked.connect(lambda: self.vugr_db())
        self.ui.action.triggered.connect(lambda: self.exit_aut())
        self.ui.action_2.triggered.connect(lambda: self.exit_personnel())
        self.write_table()
        p = session.query(patient).all()
        for i in range(session.query(patient).count()):
                self.ui.comboBox_3.addItem(str(p[i].name))

    # вывод по внешнему ключу
    def output_by_foreign_key(self,k,v,col,j):
        if k=='id_stat':
            q=session.query(user_status).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.status)))
        elif k=='id_day':
            q=session.query(day).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name)))
        elif k=='id_spec':
            q=session.query(specialization).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name_spec)))
        elif k=='id_patient':
            q=session.query(patient).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name)))
        elif k == 'id_personnel':
            q = session.query(personnel).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name)))
        elif k == 'id_reception_status':
            q = session.query(reception_status).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name)))
        elif k == 'id_office':
            q = session.query(offices).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.cab_num)))
        else:
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(v)))

    # вывод по внешнему ключу
    def output_by_foreign_key2(self, table_1, z, k, v, col):
        if k == 'id_stat':
            q = session.query(user_status).filter_by(id=int(v)).first()
            table_1.cell(z, col).text = str(q.status)
        elif k=='id_day':
            q=session.query(day).filter_by(id=int(v)).first()
            table_1.cell(z, col).text = str(q.name)
        elif k=='id_spec':
            q=session.query(specialization).filter_by(id=int(v)).first()
            table_1.cell(z, col).text = str(q.name_spec)
        elif k=='id_patient':
            q=session.query(patient).filter_by(id=int(v)).first()
            table_1.cell(z, col).text = str(q.name)
        elif k == 'id_personnel':
            q = session.query(personnel).filter_by(id=int(v)).first()
            table_1.cell(z, col).text = str(q.name)
        elif k == 'id_reception_status':
            q = session.query(reception_status).filter_by(id=int(v)).first()
            table_1.cell(z, col).text = str(q.name)
        elif k == 'id_office':
            q = session.query(offices).filter_by(id=int(v)).first()
            table_1.cell(z, col).text = str(q.cab_num)
        else:
            table_1.cell(z, col).text = str(v)

# Выгрузка БД в Docx или Xlsx
    def vugr_db(self):
        if self.ui.comboBox_5.currentText() == "В Docx":
            mydoc = docx.Document()
            for i in range(3):
                table_1 = mydoc.add_table(rows=0, cols=0, style='Table Grid')
                table = user_status
                collums = []
                collums_2 = []
                if i == 1 and (self.ui.comboBox_4.currentText() == "Все таблицы"
                               or self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной"
                             or self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом" or
                             ((self.ui.comboBox_4.currentText() == "Всю таблицу"
                               or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной"
                               or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом")
                              and self.ui.comboBox.currentText() == "Мед книжка")):
                    mydoc.add_paragraph("Таблица мед книжки (med_knigа):")
                    if self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной" \
                            or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной":
                        table_1 = mydoc.add_table(rows=session.query(med_knigа).filter_by(id_personnel=classes.glob_id).count() + 1, cols=4, style='Table Grid')
                    elif self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом" \
                            or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом":
                        p = session.query(patient).all()
                        p_id = p[self.ui.comboBox_3.currentIndex()].id
                        table_1 = mydoc.add_table(rows=session.query(med_knigа).filter_by(id_patient=p_id).count() + 1, cols=4, style='Table Grid')
                    else:
                        table_1 = mydoc.add_table(rows=session.query(med_knigа).count() + 1, cols=4, style='Table Grid')
                    table = med_knigа
                    collums = ['id', 'id_patient', 'id_personnel', 'id_spec', 'diagnoz']
                    collums_2 = ['ФИО пациента', 'ФИО персонала', 'Специализация', 'Диагноз']
                if i == 2 and (self.ui.comboBox_4.currentText() == "Все таблицы"
                               or self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной"
                             or self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом" or
                             ((self.ui.comboBox_4.currentText() == "Всю таблицу"
                               or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной"
                               or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом")
                              and self.ui.comboBox.currentText() == "Приём")):
                    mydoc.add_paragraph("Таблица записи на приём к врачу (reception):")
                    if self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной" \
                            or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной":
                        table_1 = mydoc.add_table(rows=session.query(reception).filter_by(id_personnel=classes.glob_id).count() + 1, cols=6, style='Table Grid')
                    elif self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом" \
                            or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом":
                        p = session.query(patient).all()
                        p_id = p[self.ui.comboBox_3.currentIndex()].id
                        table_1 = mydoc.add_table(rows=session.query(reception).filter_by(id_patient=p_id).count() + 1, cols=6, style='Table Grid')
                    else:
                        table_1 = mydoc.add_table(rows=session.query(reception).count() + 1, cols=6, style='Table Grid')
                    table = reception
                    collums = ['id', 'id_patient', 'id_office', 'id_reception_status', 'id_personnel', 'id_spec', 'date']
                    collums_2 = ["ФИО пациента", "№ кабинета", "Статуса приёма", "ФИО персонала", "Специализация персонала", "Дата"]
                line = session.query(table).count()
                self.ui.qTable = session.query(table).all()
                z = 0
                for j in collums_2:
                    table_1.cell(0, z).text = j
                    z += 1
                sort = [0] * line
                z = 0
                for j in session.query(table).all():
                    sort[z] = j.id
                    z += 1
                sort.sort()
                z = 0
                perem=0
                f = 1
                f2 = 1
                for j in sort:
                    if self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной" \
                            or self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом" \
                            or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной" \
                            or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом":
                        if f2 == 1:
                            z += 1
                    else:
                        z += 1
                    f2 = 0
                    for row, form in enumerate(self.ui.qTable):
                        col = 0
                        for c in collums:
                            for k, v in vars(form).items():
                                if c == k:
                                    if c == 'id':
                                        perem = v
                                        if self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной" \
                                                or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной":
                                            rec = session.query(table).filter_by(id=v).first()
                                            if classes.glob_id == rec.id_personnel:
                                                f = 1
                                            else:
                                                f = 0
                                        elif self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом" \
                                                or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом":
                                            p = session.query(patient).all()
                                            p_id = p[self.ui.comboBox_3.currentIndex()].id
                                            rec = session.query(table).filter_by(id=v).first()
                                            if p_id == rec.id_patient:
                                                f = 1
                                            else:
                                                f = 0
                                    if (c != 'id') and (perem == j) and (f == 1):
                                        self.output_by_foreign_key2(table_1, z, k, v, col)
                                        col += 1
                                        f2 = 1
            mydoc.save("C:/Users/voffc/OneDrive/Desktop/database.docx")
        else:
            table1 = pd.DataFrame({})
            table2 = pd.DataFrame({})
            for i in range(2):
                if i == 0 and (self.ui.comboBox_4.currentText() == "Все таблицы"
                               or self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной"
                             or self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом" or
                             ((self.ui.comboBox_4.currentText() == "Всю таблицу"
                               or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной"
                               or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом")
                              and self.ui.comboBox.currentText() == "Приём")):
                    if self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной"\
                            or self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной":
                        count = session.query(reception).filter_by(id_personnel=classes.glob_id).count()
                    elif self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом" \
                         or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом":
                        p = session.query(patient).all()
                        p_id = p[self.ui.comboBox_3.currentIndex()].id
                        count = session.query(reception).filter_by(id_patient=p_id).count()
                    else:
                        count = session.query(reception).count()
                    collum1 = [""] * count
                    collum2 = [""] * count
                    collum3 = [""] * count
                    collum4 = [""] * count
                    collum5 = [""] * count
                    collum6 = [""] * count
                    z = 0
                    if self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной" \
                            or self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной":
                        for j in session.query(reception).filter_by(id_personnel=classes.glob_id):
                            collum1[z] = (session.query(patient).filter_by(id=j.id_patient).first()).name
                            collum2[z] = (session.query(offices).filter_by(id=j.id_office).first()).cab_num
                            collum3[z] = (session.query(reception_status).filter_by(id=j.id_reception_status).first()).name
                            collum4[z] = (session.query(personnel).filter_by(id=j.id_personnel).first()).name
                            collum5[z] = (session.query(specialization).filter_by(id=j.id_spec).first()).name_spec
                            collum6[z] = j.date
                            z += 1
                    elif self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом" \
                         or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом":
                        p = session.query(patient).all()
                        p_id = p[self.ui.comboBox_3.currentIndex()].id
                        for j in session.query(reception).filter_by(id_patient=p_id):
                            collum1[z] = (session.query(patient).filter_by(id=j.id_patient).first()).name
                            collum2[z] = (session.query(offices).filter_by(id=j.id_office).first()).cab_num
                            collum3[z] = (session.query(reception_status).filter_by(id=j.id_reception_status).first()).name
                            collum4[z] = (session.query(personnel).filter_by(id=j.id_personnel).first()).name
                            collum5[z] = (session.query(specialization).filter_by(id=j.id_spec).first()).name_spec
                            collum6[z] = j.date
                            z += 1
                    else:
                        for j in session.query(reception).all():
                            collum1[z] = (session.query(patient).filter_by(id=j.id_patient).first()).name
                            collum2[z] = (session.query(offices).filter_by(id=j.id_office).first()).cab_num
                            collum3[z] = (session.query(reception_status).filter_by(id=j.id_reception_status).first()).name
                            collum4[z] = (session.query(personnel).filter_by(id=j.id_personnel).first()).name
                            collum5[z] = (session.query(specialization).filter_by(id=j.id_spec).first()).name_spec
                            collum6[z] = j.date
                            z += 1
                    table1 = pd.DataFrame({'ФИО пациента': collum1, '№ кабинета': collum2, 'Статус приёма': collum3, 'ФИО персонала': collum4, 'Специализация персонала': collum5, 'Дата': collum6})
                if i == 1 and (self.ui.comboBox_4.currentText() == "Все таблицы"
                               or self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной"
                             or self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом" or
                             ((self.ui.comboBox_4.currentText() == "Всю таблицу"
                               or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной"
                               or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом")
                              and self.ui.comboBox.currentText() == "Мед книжка")):
                    if self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной"\
                            or self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной":
                        count = session.query(med_knigа).filter_by(id_personnel=classes.glob_id).count()
                    elif self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом" \
                         or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом":
                        p = session.query(patient).all()
                        p_id = p[self.ui.comboBox_3.currentIndex()].id
                        count = session.query(med_knigа).filter_by(id_patient=p_id).count()
                    else:
                        count = session.query(med_knigа).count()
                    collum1 = [""] * count
                    collum2 = [""] * count
                    collum3 = [""] * count
                    collum4 = [""] * count
                    z = 0
                    if self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной" \
                            or self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной":
                        for j in session.query(med_knigа).filter_by(id_personnel=classes.glob_id):
                            collum1[z] = (session.query(patient).filter_by(id=j.id_patient).first()).name
                            collum2[z] = (session.query(personnel).filter_by(id=j.id_personnel).first()).name
                            collum3[z] = (session.query(specialization).filter_by(id=j.id_spec).first()).name_spec
                            collum4[z] = j.diagnoz
                    elif self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом" \
                         or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом":
                        p = session.query(patient).all()
                        p_id = p[self.ui.comboBox_3.currentIndex()].id
                        for j in session.query(med_knigа).filter_by(id_patient=p_id):
                            collum1[z] = (session.query(patient).filter_by(id=j.id_patient).first()).name
                            collum2[z] = (session.query(personnel).filter_by(id=j.id_personnel).first()).name
                            collum3[z] = (session.query(specialization).filter_by(id=j.id_spec).first()).name_spec
                            collum4[z] = j.diagnoz
                    else:
                        for j in session.query(med_knigа).all():
                            collum1[z] = (session.query(patient).filter_by(id=j.id_patient).first()).name
                            collum2[z] = (session.query(personnel).filter_by(id=j.id_personnel).first()).name
                            collum3[z] = (session.query(specialization).filter_by(id=j.id_spec).first()).name_spec
                            collum4[z] = j.diagnoz
                            z += 1
                    table2 = pd.DataFrame({'ФИО пациента': collum1, 'ФИО персонала': collum2, 'Специализация': collum3, 'Диагноз': collum4})
            if self.ui.comboBox_4.currentText() == "Все таблицы" \
                    or self.ui.comboBox_4.currentText() == "Все таблицы, но строки со мной" \
                    or self.ui.comboBox_4.currentText() == "Все таблицы, но строки с пациентом":
                salary_sheets = {'Запись на приём': table1, 'Мед книжка': table2}
            if (self.ui.comboBox_4.currentText() == "Всю таблицу" or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной"
                 or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом") and self.ui.comboBox.currentText() == "Мед книжка":
                salary_sheets = {'Мед книжка': table2}
            if (self.ui.comboBox_4.currentText() == "Всю таблицу" or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки со мной"
                    or self.ui.comboBox_4.currentText() == "Всю таблицу, но строки с пациентом") and self.ui.comboBox.currentText() == "Приём":
                salary_sheets = {'Запись на приём': table1}
            writer = pd.ExcelWriter('./database123456.xlsx', engine='xlsxwriter')
            for sheet_name in salary_sheets.keys():
                salary_sheets[sheet_name].to_excel(writer, sheet_name=sheet_name, index=False)
            writer.save()

    def write_table(self):
        self.ui.tableWidget.clearSelection()
        result = self.ui.comboBox.currentText()
        table = patient
        self.ui.tableWidget.setColumnCount(1)
        collums = ['name']
        self.ui.tableWidget.setHorizontalHeaderLabels(["ФИО"])
        if result == "Пациенты":
            table = patient
            self.ui.tableWidget.setColumnCount(1)
            collums = ['name']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ФИО"])
        if result == "Персонал":
            table = personnel
            self.ui.tableWidget.setColumnCount(3)
            collums = ['id', 'id_spec', 'id_office', 'name']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                ["Специализация персонала", "№ кабинета", "ФИО"])
        if result == "Время работы":
            table = working_hours
            self.ui.tableWidget.setColumnCount(5)
            collums = ['id', 'id_day', 'id_spec', 'work_hours', 'free_time', 'break_time']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                [ "День", "Специализация", "Рабочее время", "Свободное время", "Перерыв"])
        if result == "Приём":
            table = reception
            self.ui.tableWidget.setColumnCount(6)
            collums = ['id', 'id_patient', 'id_office', 'id_reception_status', 'id_personnel', 'id_spec', 'date']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                ["ФИО пациента", "№ кабинета", "Статуса приёма", "ФИО персонала", "Специализация персонала", "Дата"])
        if result == "Мед книжка":
            table = med_knigа
            self.ui.tableWidget.setColumnCount(4)
            collums = ['id', 'id_patient', 'id_personnel', 'id_spec', 'diagnoz']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ФИО пациента", "ФИО персонала", 'Специализация', "Диагноз"])
        line = session.query(table).count()
        self.ui.qTable = session.query(table).all()
        self.ui.tableWidget.setRowCount(line)
        sort = [0] * line
        j = 0
        for row, form in enumerate(self.ui.qTable):
            col = 0
            for c in collums:
                for k, v in vars(form).items():
                    if c == k:
                        self.ui.tableWidget.setItem(row, col, QTableWidgetItem(str(v)))
                        col += 1
        for i in session.query(table).all():
            sort[j] = i.id
            j += 1
        sort.sort()
        perem = 0
        j = -1
        for i in sort:
            j += 1
            for row, form in enumerate(self.ui.qTable):
                col = 0
                for c in collums:
                    for k, v in vars(form).items():
                        if c == k:
                            if c == 'id':
                                perem = v
                            if (c != 'id') & (perem == i):
                                self.output_by_foreign_key(k,v,col,j)
                                col += 1
        self.ui.tableWidget.resizeColumnsToContents()

    # Выход из окна персонала обратно в окно авторизации
    def exit_aut(self):
        self.hide()
        dialog = log_window.log_panel(parent=self)
        dialog.show()

    # Выход из окна выгрузки обратно в окно персонала
    def exit_personnel(self):
        self.hide()
        dialog = personnel_window.personal_panel(parent=self)
        dialog.show()