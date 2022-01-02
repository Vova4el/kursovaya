from Qt import *
from backup_bd import *
from Authentication import login_panell
import classes
from personalQT import personalUI
from patientQT import patientUI
from personal_med_kniga_QT import Med_knigaUi
from datetime import timedelta, datetime, time
import docx
import pandas as pd
#ОКНО_АДМИНА
class admin_panel(QMainWindow):
    def __init__(self, parent=None):
        super(admin_panel, self).__init__(parent)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(lambda: self.write_table())
        self.ui.pushButton_2.clicked.connect(lambda: self.change_table())
        self.ui.pushButton_3.clicked.connect(lambda: self.confirm_change())
        self.ui.pushButton_4.clicked.connect(lambda: self.undo_change())
        self.ui.pushButton_5.clicked.connect(lambda: self.vugr_db())
        self.ui.comboBox_5.addItem("В Docx")
        self.ui.comboBox_5.addItem("В Xlsx")

    # Выгрузка БД в Docx или Xlsx
    def vugr_db(self):
        if self.ui.comboBox_5.currentText() == "В Docx":
            mydoc = docx.Document()
            for i in range(12):
                table_1 = mydoc.add_table(rows=0, cols=0, style='Table Grid')
                table = user_status
                collums = []
                if i == 1:
                    mydoc.add_paragraph("Таблица статусов пользователей (user_status):")
                    table_1 = mydoc.add_table(rows=session.query(user_status).count() + 1, cols=2, style='Table Grid')
                    table = user_status
                    collums = ['id', 'status']
                if i == 2:
                    mydoc.add_paragraph("Таблица пациенитов (patient):")
                    table_1 = mydoc.add_table(rows=session.query(patient).count() + 1, cols=3, style='Table Grid')
                    table = patient
                    collums = ['id', 'name', 'id_acc']
                if i == 3:
                    mydoc.add_paragraph("Таблица мед книжки (med_knigа):")
                    table_1 = mydoc.add_table(rows=session.query(med_knigа).count() + 1, cols=4, style='Table Grid')
                    table = med_knigа
                    collums = ['id', 'id_patient', 'id_personnel', 'diagnoz']
                if i == 4:
                    mydoc.add_paragraph("Таблица дней (day):")
                    table_1 = mydoc.add_table(rows=session.query(day).count() + 1, cols=2, style='Table Grid')
                    table = day
                    collums = ['id', 'name']
                if i == 5:
                    mydoc.add_paragraph("Таблица специализаций (specialization):")
                    table_1 = mydoc.add_table(rows=session.query(specialization).count() + 1, cols=2, style='Table Grid')
                    table = specialization
                    collums = ['id', 'name_spec']
                if i == 6:
                    mydoc.add_paragraph("Таблица аккаунтов (accounts):")
                    table_1 = mydoc.add_table(rows=session.query(accounts).count() + 1, cols=4, style='Table Grid')
                    table = accounts
                    collums = ['id', 'login', 'password', 'id_stat']
                if i == 7:
                    mydoc.add_paragraph("Таблица статуса приёма (reception_status):")
                    table_1 = mydoc.add_table(rows=session.query(reception_status).count() + 1, cols=2, style='Table Grid')
                    table = reception_status
                    collums = ['id', 'name']
                if i == 8:
                    mydoc.add_paragraph("Таблица кабинетов (offices):")
                    table_1 = mydoc.add_table(rows=session.query(offices).count() + 1, cols=2, style='Table Grid')
                    table = offices
                    collums = ['id', 'cab_num']
                if i == 9:
                    mydoc.add_paragraph("Рабочее время (working_hours):")
                    table_1 = mydoc.add_table(rows=session.query(working_hours).count() + 1, cols=6, style='Table Grid')
                    table = working_hours
                    collums = ['id', 'id_day', 'id_spec', 'work_hours', 'free_time', 'break_time']
                if i == 10:
                    mydoc.add_paragraph("Таблица записи на приём к врачу (reception):")
                    table_1 = mydoc.add_table(rows=session.query(reception).count() + 1, cols=7, style='Table Grid')
                    table = reception
                    collums = ['id', 'id_patient', 'id_office', 'id_reception_status', 'id_personnel', 'id_spec','date']
                if i == 11:
                    mydoc.add_paragraph("Таблица записи на приём к врачу (personnel):")
                    table_1 = mydoc.add_table(rows=session.query(personnel).count() + 1, cols=5, style='Table Grid')
                    table = personnel
                    collums = ['id', 'id_spec', 'id_office', 'name', 'id_acc']
                line = session.query(table).count()
                self.ui.qTable = session.query(table).all()
                z = 0
                for j in collums:
                    table_1.cell(0, z).text = j
                    z += 1
                sort = [0] * line
                z = 0
                for j in session.query(table).all():
                    sort[z] = j.id
                    z += 1
                sort.sort()
                perem = 0
                z = 0
                for j in sort:
                    z += 1
                    for row, form in enumerate(self.ui.qTable):
                        col = 0
                        for c in collums:
                            for k, v in vars(form).items():
                                if c == k:
                                    if c == 'id':
                                        perem = v
                                    if ((c == 'id') & (v == j)) | ((c != 'id') & (perem == j)):
                                        table_1.cell(z, col).text = str(v)
                                        col += 1

            mydoc.save("C:/Users/voffc/OneDrive/Desktop/database.docx")
        else:
            table1 = pd.DataFrame({})
            table2 = pd.DataFrame({})
            table3 = pd.DataFrame({})
            table4 = pd.DataFrame({})
            table5 = pd.DataFrame({})
            table6 = pd.DataFrame({})
            table7 = pd.DataFrame({})
            table8 = pd.DataFrame({})
            table9 = pd.DataFrame({})
            table10 = pd.DataFrame({})
            table11 = pd.DataFrame({})
            for i in range(11):
                if i == 0:
                    count = session.query(personnel).count()
                    collum1 = [0] * count
                    collum2 = [0] * count
                    collum3 = [0] * count
                    collum4 = [""] * count
                    collum5 = [0] * count
                    z = 0
                    for j in session.query(personnel).all():
                        collum1[z] = j.id
                        collum2[z] = j.id_spec
                        collum3[z] = j.id_office
                        collum4[z] = j.name
                        collum5[z] = j.id_acc
                        z += 1
                    table1 = pd.DataFrame({'id': collum1, 'id_spec': collum2, 'id_office': collum3, 'name': collum4, 'id_acc':collum5})
                if i == 1:
                    count = session.query(patient).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    collum3 = [0] * count
                    z = 0
                    for j in session.query(patient).all():
                        collum1[z] = j.id
                        collum2[z] = j.name
                        collum3[z] = j.id_acc
                        z += 1
                    table2 = pd.DataFrame({'id': collum1, 'name': collum2,'id_acc': collum3})
                if i == 2:
                    count = session.query(user_status).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    z = 0
                    for j in session.query(user_status).all():
                        collum1[z] = j.id
                        collum2[z] = j.status
                        z += 1
                    table3 = pd.DataFrame({'id': collum1, 'status': collum2})
                if i == 3:
                    count = session.query(offices).count()
                    collum1 = [0] * count
                    collum2 = [0] * count
                    z = 0
                    for j in session.query(offices).all():
                        collum1[z] = j.id
                        collum2[z] = j.cab_num
                        z += 1
                    table4 = pd.DataFrame({'id': collum1, 'cab_num': collum2})
                if i == 4:
                    count = session.query(working_hours).count()
                    collum1 = [0] * count
                    collum2 = [0] * count
                    collum3 = [0] * count
                    collum4 = [""] * count
                    collum5 = [""] * count
                    collum6 = [""] * count
                    z = 0
                    for j in session.query(working_hours).all():
                        collum1[z] = j.id
                        collum2[z] = j.id_day
                        collum3[z] = j.id_spec
                        collum4[z] = j.work_hours
                        collum5[z] = j.free_time
                        collum6[z] = j.break_time
                        z += 1
                    table5 = pd.DataFrame({'id': collum1, 'id_day': collum2, 'id_spec': collum3, 'work_hours': collum4,
                                           'free_time': collum5, 'break_time': collum6})
                if i == 5:
                    count = session.query(reception).count()
                    collum1 = [0] * count
                    collum2 = [0] * count
                    collum3 = [0] * count
                    collum4 = [0] * count
                    collum5 = [0] * count
                    collum6 = [0] * count
                    collum7 = [""] * count
                    z = 0
                    for j in session.query(reception).all():
                        collum1[z] = j.id
                        collum2[z] = j.id_patient
                        collum3[z] = j.id_office
                        collum4[z] = j.id_reception_status
                        collum5[z] = j.id_personnel
                        collum6[z] = j.id_spec
                        collum7[z] = j.date
                        z += 1
                    table6 = pd.DataFrame({'id': collum1, 'id_patient': collum2, 'id_office': collum3, 'id_reception_status': collum4, 'id_personnel': collum5, 'id_spec': collum6, 'date':collum7})
                if i == 6:
                    count = session.query(accounts).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    collum3 = [""] * count
                    collum4 = [0] * count
                    z = 0
                    for j in session.query(accounts).all():
                        collum1[z] = j.id
                        collum2[z] = j.login
                        collum3[z] = j.password
                        collum4[z] = j.id_stat
                        z += 1
                    table7 = pd.DataFrame({'id': collum1, 'login': collum2, 'password': collum3, 'id_stat': collum4})
                if i == 7:
                    count = session.query(med_knigа).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    collum3 = [0] * count
                    collum4 = [""] * count
                    z = 0
                    for j in session.query(med_knigа).all():
                        collum1[z] = j.id
                        collum2[z] = j.id_patient
                        collum3[z] = j.id_personnel
                        collum4[z] = j.diagnoz
                        z += 1
                    table8 = pd.DataFrame({'id': collum1, 'id_patient': collum2, 'id_personnel': collum3,
                                           'diagnoz': collum4})
                if i == 8:
                    count = session.query(day).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    z = 0
                    for j in session.query(day).all():
                        collum1[z] = j.id
                        collum2[z] = j.name
                        z += 1
                    table9 = pd.DataFrame({'id': collum1, 'name': collum2})
                if i == 9:
                    count = session.query(specialization).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    z = 0
                    for j in session.query(specialization).all():
                        collum1[z] = j.id
                        collum2[z] = j.name_spec
                        z += 1
                    table10 = pd.DataFrame({'id': collum1, 'name_spec': collum2})
                if i == 10:
                    count = session.query(reception_status).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    z = 0
                    for j in session.query(reception_status).all():
                        collum1[z] = j.id
                        collum2[z] = j.name
                        z += 1
                    table11 = pd.DataFrame({'id': collum1, 'name': collum2})
            salary_sheets = {'Персонал': table1, 'Пациенты': table2, 'Статусы пользователей': table3,
                             'Кабинеты': table4, 'Рабочее время': table5, 'Запись на приём': table6,
                             'Аккаунты': table7, 'Мед книжка': table8, 'Дни': table9, 'Специальность': table10,
                             'Статусы записей на приём': table11}
            writer = pd.ExcelWriter('./database123456.xlsx', engine='xlsxwriter')
            for sheet_name in salary_sheets.keys():
                salary_sheets[sheet_name].to_excel(writer, sheet_name=sheet_name, index=False)
            writer.save()

    def write_table(self):
        self.ui.comboBox_2.clear()
        self.ui.comboBox_3.clear()
        self.ui.tableWidget.clearSelection()
        result = self.ui.comboBox.currentText()
        collums = []
        if result == "Аккаунты":
            table = accounts
            self.ui.tableWidget.setColumnCount(4)
            collums = ['id', 'login', 'password', 'id_stat']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID", "Логин", "Пароль", "ID статуса"])
        if result == "Статусы пользователей":
            table = user_status
            self.ui.tableWidget.setColumnCount(2)
            collums = ['id', 'status']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID", "Cтатус"])
        if result == "Пациенты":
            table = patient
            self.ui.tableWidget.setColumnCount(3)
            collums = ['id', 'name', 'id_acc']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID", "ФИО", "ID аккаунта"])
        if result == "День":
            table = day
            self.ui.tableWidget.setColumnCount(2)
            collums = ['id', 'name']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                ["ID", "День"])
        if result == "Специализация":
            table = specialization
            self.ui.tableWidget.setColumnCount(2)
            collums = ['id', 'name_spec']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID", "Специализация"])
        if result == "Статус приёма":
            table = reception_status
            self.ui.tableWidget.setColumnCount(2)
            collums = ['id', 'name']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID", "Статус"])
        if result == "Кабинеты":
            table = offices
            self.ui.tableWidget.setColumnCount(2)
            collums = ['id', 'cab_num']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID", "Номер кабинета"])
        if result == "Персонал":
            table = personnel
            self.ui.tableWidget.setColumnCount(5)
            collums = ['id', 'id_spec', 'id_office', 'name', 'id_acc']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                ["ID", "ID специализации", "ID кабинета", "ФИО", 'ID аккаунта'])
        if result == "Время работы":
            table = working_hours
            self.ui.tableWidget.setColumnCount(6)
            collums = ['id', 'id_day', 'id_spec', 'work_hours', 'free_time', 'break_time']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                ["ID", "ID дня", "ID специализации", "рабочее время", "свободное время", "перерыв"])
        if result == "Приём":
            table = reception
            self.ui.tableWidget.setColumnCount(7)
            collums = ['id', 'id_patient', 'id_office', 'id_reception_status', 'id_personnel', 'id_spec', 'date']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                ["ID", "ID пациента", "ID кабинета", "ID статуса приёма", "ID персонала", "ID специализации", "Дата"])
        if result == "Мед книжка":
            table = med_knigа
            self.ui.tableWidget.setColumnCount(4)
            collums = ['id', 'id_patient','id_personnel', 'diagnoz']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID", "ID пациента", "ID персонала", "Диагноз"])
        line = session.query(table).count()
        self.ui.qTable = session.query(table).all()
        self.ui.tableWidget.setRowCount(line)
        sort = [0] * line
        j = 0
        for i in collums:
            if i != 'id':
                self.ui.comboBox_2.addItem(i)
        for i in session.query(table).all():
            sort[j] = i.id
            j += 1
        sort.sort()
        for i in sort:
            self.ui.comboBox_3.addItem(str(i))
        perem = 0
        j = -1
        for i in sort:
            j += 1
            for row, form in enumerate(self.ui.qTable):
                col = 0
                for c in collums:
                    for k, v in vars(form).items():
                        if c == k:
                            if c == 'id':
                                perem = v
                            if ((c == 'id') & (v == i)) | ((c != 'id') & (perem == i)):
                                self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(v)))
                                col += 1
        self.ui.tableWidget.setStyleSheet("selection-color: rgb(255, 0, 127);\n"
                                       "selection-background-color: rgb(85, 255, 127);")
        self.ui.tableWidget.resizeColumnsToContents()


    def change_table(self):
        self.ui.statusbar.clearMessage()
        if self.ui.prov == 0:
            self.ui.prov = 1
        result = self.ui.comboBox.currentText()
        change = self.ui.comboBox_4.currentText()
        colona = self.ui.comboBox_2.currentText()
        stlb = int(self.ui.comboBox_3.currentText())
        text = self.ui.textEdit.toPlainText()
        lis = text.split(',')
        table = patient
        try:
            if result == "Аккаунты":
                table = accounts
                if change == "Обновить элемент":
                    query = session.query(table).get(stlb)
                    if colona == "login":
                        query.login = text
                    if colona == "password":
                        query.password = text
                    if colona == "id_stat":
                        query.id_stat = int(text)
                if change == "Добавить строку":
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    f = 0  # флаг
                    free_id = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id = free_id + 1
                            f = 0
                        else:
                            f = 1
                    new = table(id=free_id, login=lis[0], password=lis[1], id_stat=int(lis[2]))
                    session.add(new)
            if result == "Статусы пользователей":
                table = user_status
                if change == "Обновить элемент":
                    query = session.query(table).get(stlb)
                    if colona == "status":
                        query.status = text
                if change == "Добавить строку":
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    f = 0  # флаг
                    free_id = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id = free_id + 1
                            f = 0
                        else:
                            f = 1
                    new = table(id=free_id, status=text)
                    session.add(new)
            if result == "Пациенты":
                table = patient
                if change == "Обновить элемент":
                    query = session.query(table).get(stlb)
                    if colona == "name":
                        query.name = text
                    if colona == "id_acc":
                        query.id_acc = int(text)
                if change == "Добавить строку":
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    f = 0  # флаг
                    free_id = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id = free_id + 1
                            f = 0
                        else:
                            f = 1
                    new = table(id=free_id, name=lis[0],id_acc = int(lis[1]))
                    session.add(new)
            if result == "День":
                table = day
                if change == "Обновить элемент":
                    query = session.query(table).get(stlb)
                    if colona == "name":
                        query.name = text
                if change == "Добавить строку":
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    f = 0  # флаг
                    free_id = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id = free_id + 1
                            f = 0
                        else:
                            f = 1
                    new = table(id=free_id, name=text)
                    session.add(new)
            if result == "Специализация":
                table = specialization
                if change == "Обновить элемент":
                    query = session.query(table).get(stlb)
                    if colona == "name_spec":
                        query.name_spec = text
                if change == "Добавить строку":
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    f = 0  # флаг
                    free_id = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id = free_id + 1
                            f = 0
                        else:
                            f = 1
                    new = table(id=free_id, name_spec=text)
                    session.add(new)
            if result == "Статус приёма":
                table = reception_status
                if change == "Обновить элемент":
                    query = session.query(table).get(stlb)
                    if colona == "name":
                        query.name = text
                if change == "Добавить строку":
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    f = 0  # флаг
                    free_id = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id = free_id + 1
                            f = 0
                        else:
                            f = 1
                    new = table(id=free_id, name=text)
                    session.add(new)
            if result == "Кабинеты":
                table = offices
                if change == "Обновить элемент":
                    query = session.query(table).get(stlb)
                    if colona == "cab_num":
                        query.cab_num = int(text)
                if change == "Добавить строку":
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    f = 0  # флаг
                    free_id = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id = free_id + 1
                            f = 0
                        else:
                            f = 1
                    new = table(id=free_id, cab_num=int(text))
                    session.add(new)
            if result == "Персонал":
                table = personnel
                if change == "Обновить элемент":
                    query = session.query(table).get(stlb)
                    if colona == "id_spec":
                        query.id_spec = int(text)
                    if colona == "id_office":
                        query.id_office = int(text)
                    if colona == "name":
                        query.name = text
                    if colona == "id_acc":
                        query.id_acc = int(text)
                if change == "Добавить строку":
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    f = 0  # флаг
                    free_id = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id = free_id + 1
                            f = 0
                        else:
                            f = 1
                    new = table(id=free_id, id_spec=int(lis[0]), id_office=int(lis[1]), name= lis[2], id_acc=int(lis[3]))
                    session.add(new)
            if result == "Время работы":
                table = working_hours
                if change == "Обновить элемент":
                    query = session.query(table).get(stlb)
                    if colona == "id_day":
                        query.id_day = int(text)
                    if colona == "work_hours":
                        query.work_hours= [lis[0],lis[1]]
                    if colona == "free_time":
                        query.free_time = [lis[0],lis[1]]
                    if colona == "break_time":
                        query.break_time = [lis[0],lis[1]]
                if change == "Добавить строку":
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    f = 0  # флаг
                    free_id = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id = free_id + 1
                            f = 0
                        else:
                            f = 1
                    new = table(id=free_id, id_day=int(lis[0]), id_spec=int(lis[1]), work_hours= [lis[2],lis[3]], free_time=[lis[4],lis[5]], break_time=[lis[6],lis[7]])
                    session.add(new)
            if result == "Приём":
                table = reception
                if change == "Обновить элемент":
                    query = session.query(table).get(stlb)
                    if colona == "id_patient":
                        query.id_patient = int(text)
                    if colona == "id_office":
                        query.id_office= int(text)
                    if colona == "id_reception_status":
                        query.id_reception_status = int(text)
                    if colona == "id_personnel":
                        query.id_personnel = int(text)
                    if colona == "id_spec":
                        query.id_spec = int(text)
                    if colona == "date":
                        query.date = text
                if change == "Добавить строку":
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    f = 0  # флаг
                    free_id = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id = free_id + 1
                            f = 0
                        else:
                            f = 1
                    new = table(id=free_id, id_patient=int(lis[0]), id_office=int(lis[1]), id_reception_status= int(lis[2]), id_personnel=int(lis[3]), id_spec=int(lis[4]), date=lis[5]) #1,2,2,1,1,2006-12-14 20:15:00
                    session.add(new)
            if result == "Мед книжка":
                table = med_knigа
                if change == "Обновить элемент":
                    query = session.query(table).get(stlb)
                    if colona == "id_patient":
                        query.id_patient = int(text)
                    if colona == "id_personnel":
                        query.id_personnel= int(text)
                    if colona == "diagnoz":
                        query.diagnoz = text
                if change == "Добавить строку":
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    f = 0  # флаг
                    free_id = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id = free_id + 1
                            f = 0
                        else:
                            f = 1
                    new = table(id=free_id, id_patient=int(lis[0]), id_personnel=int(lis[1]), diagnoz=lis[2])#id,id_patient, id_personnel, diagnoz
                    session.add(new)
            if change == "Удалить строку":
                session.query(table).filter_by(id=stlb).delete(synchronize_session=False)
            self.write_table()

        except:
            self.ui.statusbar.showMessage("Произошла ошибка")

    def confirm_change(self):
        if self.ui.prov != 0:
            session.commit()
            dump_sqlalchemy()
            self.ui.prov = 0
            self.write_table()
        else:
            self.ui.statusbar.showMessage("Сначала внесите изменение в таблицу")

    def undo_change(self):
        if self.ui.prov != 0:
            session.rollback()
            self.ui.prov = 0
            self.write_table()
        else:
            self.ui.statusbar.showMessage("Сначала внесите изменение в таблицу")

#ОКНО_ДЛЯ_ПЕРСОНАЛА#
class personal_panel(QMainWindow):
    def __init__(self, parent=None):
        super(personal_panel, self).__init__(parent)
        self.ui = personalUI()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(lambda: self.write_table())
        self.ui.pushButton_2.clicked.connect(lambda: self.exit_med_kniga())
        self.ui.action.triggered.connect(lambda: self.exit_db_panel())

    # вывод по внешнему ключу
    def output_by_foreign_key(self,k,v,col,j):
        if k=='id_stat':
            q=session.query(user_status).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.status)))
        elif k=='id_day':
            q=session.query(day).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name)))
        elif k=='id_spec':
            q=session.query(specialization).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name_spec)))
        elif k=='id_patient':
            q=session.query(patient).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name)))
        elif k == 'id_personnel':
            q = session.query(personnel).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name)))
        elif k == 'id_reception_status':
            q = session.query(reception_status).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name)))
        elif k == 'id_office':
            q = session.query(offices).filter_by(id=int(v)).first()
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.cab_num)))
        else:
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(v)))

    def write_table(self):
        self.ui.tableWidget.clearSelection()
        result = self.ui.comboBox.currentText()
        line = 0
        collums = []
        if result == "Пациенты":
            table = patient
            self.ui.tableWidget.setColumnCount(1)
            collums = ['name']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ФИО"])
        if result == "Персонал":
            table = personnel
            self.ui.tableWidget.setColumnCount(3)
            collums = ['id', 'id_spec', 'id_office', 'name']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                ["Специализация персонала", "№ кабинета", "ФИО"])
        if result == "Время работы":
            table = working_hours
            self.ui.tableWidget.setColumnCount(5)
            collums = ['id', 'id_day', 'id_spec', 'work_hours', 'free_time', 'break_time']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                [ "День", "Специализация", "Рабочее время", "Свободное время", "Перерыв"])
        if result == "Приём":
            table = reception
            self.ui.tableWidget.setColumnCount(6)
            collums = ['id', 'id_patient', 'id_office', 'id_reception_status', 'id_personnel', 'id_spec', 'date']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                ["ФИО пациента", "№ кабинета", "Статуса приёма", "ФИО персонала", "Специализация персонала", "Дата"])
        if result == "Мед книжка":
            table = med_knigа
            self.ui.tableWidget.setColumnCount(3)
            collums = ['id', 'id_patient', 'id_personnel', 'diagnoz']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID пациента", "ID персонала", "Диагноз"])
        line = session.query(table).count()
        self.ui.qTable = session.query(table).all()
        self.ui.tableWidget.setRowCount(line)
        sort = [0] * line
        j = 0
        for row, form in enumerate(self.ui.qTable):
            col = 0
            for c in collums:
                for k, v in vars(form).items():
                    if c == k:
                        self.ui.tableWidget.setItem(row, col, QTableWidgetItem(str(v)))
                        col += 1
        for i in session.query(table).all():
            sort[j] = i.id
            j += 1
        sort.sort()
        perem = 0
        j = -1
        for i in sort:
            j += 1
            for row, form in enumerate(self.ui.qTable):
                col = 0
                for c in collums:
                    for k, v in vars(form).items():
                        if c == k:
                            if c == 'id':
                                perem = v
                            if (c != 'id') & (perem == i):
                                self.output_by_foreign_key(k,v,col,j)
                                col += 1
        self.ui.tableWidget.resizeColumnsToContents()

    def change_table(self):
        self.ui.statusbar.clearMessage()
        if self.ui.prov == 0:
            self.ui.prov = 1
        change = self.ui.comboBox_4.currentText()
        stlb = int(self.ui.comboBox_3.currentText())
        p_id = self.ui.comboBox_2.currentText()
        text = self.ui.textEdit.toPlainText()
        try:
            if change == "Добавить строку":
                line = session.query(med_knigа).count()
                tabl = session.query(med_knigа).all()
                f = 0  # флаг
                free_id = 1  # свободный индекс
                while f == 0:
                    for i in range(line):
                        if free_id == tabl[i].id:
                            f = 1
                    if f == 1:
                        free_id = free_id + 1
                        f = 0
                    else:
                        f = 1
                new = med_knigа(id=free_id, id_patient=p_id,id_personnel=classes.glob_id,diagnoz=text)
                session.add(new)
            if change == "Обновить элемент":
                query = session.query(med_knigа).get(stlb)
                if classes.glob_id == query.id_personnel:
                    query.diagnoz = text
                else:
                    self.ui.statusbar.showMessage("Вы должны быть той же специальности, что и врач, написавший этот диагноз")
            if change == "Удалить строку":
                query = session.query(med_knigа).get(stlb)
                if classes.glob_id == query.id_personnel:
                    session.query(med_knigа).filter_by(id=stlb).delete(synchronize_session=False)
                else:
                    self.ui.statusbar.showMessage("Вы должны быть той же специальности, что и врач, написавший этот диагноз")
            self.write_table()

        except:
            self.ui.statusbar.showMessage("Произошла ошибка")

    # Выход из окна персонала обратно в окно авторизации
    def exit_db_panel(self):
        print(23)
        self.hide()
        dialog = log_panel(parent=self)
        dialog.show()

    # Выход из окна персонала обратно в окно авторизации
    def exit_med_kniga(self):
        print(23)
        self.hide()
        print(24)
        dialog = panel_med_kniga(parent=self)
        print(25)
        dialog.show()

    def confirm_change(self):
        if self.ui.prov != 0:
            session.commit()
            dump_sqlalchemy()
            self.ui.prov = 0
            self.write_table()
        else:
            self.ui.statusbar.showMessage("Сначала внесите изменение в таблицу")

    def undo_change(self):
        if self.ui.prov != 0:
            session.rollback()
            self.ui.prov = 0
            self.write_table()
        else:
            self.ui.statusbar.showMessage("Сначала внесите изменение в таблицу")




#########################Окно_ДЛЯ_работы с мед книжкой##########################################
class panel_med_kniga(QMainWindow):
    def __init__(self, parent=None):
        super(panel_med_kniga, self).__init__(parent)
        self.ui = Med_knigaUi()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(lambda: self.write_table())
        self.ui.pushButton_2.clicked.connect(lambda: self.change_table())
        self.ui.action.triggered.connect(lambda: self.exit_aut())
        self.ui.action_2.triggered.connect(lambda: self.exit_personnel())

    # вывод по внешнему ключу
    def output_by_foreign_key(self,k,v,col,j):
        print('j')
        print(f'kk=',k)
        if k=='id_stat':
            q=session.query(user_status).filter_by(id=int(v)).first()
            print(f'q',q)
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.status)))
        elif k=='id_day':
            q=session.query(day).filter_by(id=int(v)).first()
            print(f'q',q)
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name)))
        elif k=='id_spec':
            q=session.query(specialization).filter_by(id=int(v)).first()
            print(f'q',q)
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name_spec)))
        elif k=='id_patient':
            q=session.query(patient).filter_by(id=int(v)).first()
            print(f'q',q)
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name)))
        elif k == 'id_personnel':
            q = session.query(personnel).filter_by(id=int(v)).first()
            print(f'q', q)
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name)))
        elif k == 'id_reception_status':
            q = session.query(reception_status).filter_by(id=int(v)).first()
            print(f'q', q)
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.name)))
        elif k == 'id_office':
            q = session.query(offices).filter_by(id=int(v)).first()
            print(f'q', q)
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(q.cab_num)))
        else:
            self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(v)))

    def write_table(self):
        self.ui.comboBox_2.clear()
        self.ui.comboBox_3.clear()
        self.ui.tableWidget.clearSelection()
        result = self.ui.comboBox.currentText()
        line = 0
        collums = []
        table = med_knigа
        self.ui.tableWidget.setColumnCount(3)
        collums = ['id', 'id_patient', 'id_personnel', 'diagnoz']
        self.ui.tableWidget.setHorizontalHeaderLabels(["ID пациента", "ID персонала", "Диагноз"])
        line = session.query(table).count()
        self.ui.qTable = session.query(table).all()
        self.ui.tableWidget.setRowCount(line)
        sort = [0] * line
        j = 0
        p = session.query(patient).all()
        for i in range(session.query(patient).count()):
                self.ui.comboBox_2.addItem(str(p[i].name))
        for i in range(line):
            self.ui.comboBox_3.addItem(str(i + 1))
        for row, form in enumerate(self.ui.qTable):
            col = 0
            for c in collums:
                for k, v in vars(form).items():
                    if c == k:
                        self.ui.tableWidget.setItem(row, col, QTableWidgetItem(str(v)))
                        col += 1
        for i in session.query(table).all():
            sort[j] = i.id
            j += 1
        perem = 0
        j = -1
        for i in sort:
            j += 1
            for row, form in enumerate(self.ui.qTable):
                col = 0
                for c in collums:
                    for k, v in vars(form).items():
                        if c == k:
                            if c == 'id':
                                perem = v
                            if (c != 'id') & (perem == i):
                                self.output_by_foreign_key(k,v,col,j)
                                col += 1
        self.ui.tableWidget.resizeColumnsToContents()

    def change_table(self):
        self.ui.statusbar.clearMessage()
        if self.ui.prov == 0:
            self.ui.prov = 1
        change = self.ui.comboBox_4.currentText()
        stlb = int(self.ui.comboBox_3.currentText())
        p_id = self.ui.comboBox_2.currentText()
        text = self.ui.textEdit.toPlainText()
        try:
            if change == "Добавить строку":
                line = session.query(med_knigа).count()
                tabl = session.query(med_knigа).all()
                f = 0  # флаг
                free_id = 1  # свободный индекс
                while f == 0:
                    for i in range(line):
                        if free_id == tabl[i].id:
                            f = 1
                    if f == 1:
                        free_id = free_id + 1
                        f = 0
                    else:
                        f = 1
                new = med_knigа(id=free_id, id_patient=p_id,id_personnel=classes.glob_id,diagnoz=text)
                session.add(new)
            if change == "Обновить элемент":
                query = session.query(med_knigа).get(stlb)
                if classes.glob_id == query.id_personnel:
                    query.diagnoz = text
                else:
                    self.ui.statusbar.showMessage("Вы должны быть той же специальности, что и врач, написавший этот диагноз")
            if change == "Удалить строку":
                query = session.query(med_knigа).get(stlb)
                if classes.glob_id == query.id_personnel:
                    session.query(med_knigа).filter_by(id=stlb).delete(synchronize_session=False)
                else:
                    self.ui.statusbar.showMessage("Вы должны быть той же специальности, что и врач, написавший этот диагноз")
            self.write_table()

        except:
            self.ui.statusbar.showMessage("Произошла ошибка")

    # Выход из окна персонала обратно в окно авторизации
    def exit_aut(self):
        print(23)
        self.hide()
        dialog = log_panel(parent=self)
        dialog.show()

    # Выход из окна персонала обратно в окно авторизации
    def exit_personnel(self):
        print(23)
        self.hide()
        dialog = personal_panel(parent=self)
        dialog.show()

    def confirm_change(self):
        if self.ui.prov != 0:
            session.commit()
            dump_sqlalchemy()
            self.ui.prov = 0
            self.write_table()
        else:
            self.ui.statusbar.showMessage("Сначала внесите изменение в таблицу")

    def undo_change(self):
        if self.ui.prov != 0:
            session.rollback()
            self.ui.prov = 0
            self.write_table()
        else:
            self.ui.statusbar.showMessage("Сначала внесите изменение в таблицу")




#################################ОКНО_ПАЦИЕНТА######################################
class patient_panel(QMainWindow):
    def __init__(self, parent=None):
        super(patient_panel, self).__init__(parent)
        self.ui = patientUI()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(lambda: self.write_table())
        self.ui.pushButton_2.clicked.connect(lambda: self.change_table())
        self.ui.pushButton_3.clicked.connect(lambda: self.confirm_change())
        self.ui.pushButton_4.clicked.connect(lambda: self.undo_change())
        self.ui.pushButton_5.clicked.connect(lambda: self.autoplay_spec())


    def write_table(self):
        self.ui.comboBox_3.clear()
        self.ui.comboBox_5.clear()
        self.ui.tableWidget.clearSelection()
        result = self.ui.comboBox.currentText()
        line = 0
        collums = []
        if result == "Пациенты":
            table = patient
            self.ui.tableWidget.setColumnCount(3)
            collums = ['id', 'name', 'id_acc']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID", "ФИО", "ID аккаунта"])
        if result == "День":
            table = day
            self.ui.tableWidget.setColumnCount(2)
            collums = ['id', 'name']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                ["ID", "День"])
        if result == "Специализация":
            table = specialization
            self.ui.tableWidget.setColumnCount(2)
            collums = ['id', 'name_spec']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID", "Специализация"])
        if result == "Статус приёма":
            table = reception_status
            self.ui.tableWidget.setColumnCount(2)
            collums = ['id', 'name']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID", "Статус"])
        if result == "Кабинеты":
            table = offices
            self.ui.tableWidget.setColumnCount(2)
            collums = ['id', 'cab_num']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID", "Номер кабинета"])
        if result == "Персонал":
            table = personnel
            self.ui.tableWidget.setColumnCount(5)
            collums = ['id', 'id_spec', 'id_office', 'name', 'id_acc']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                ["ID", "ID специализации", "ID кабинета", "ФИО", 'ID аккаунта'])
        if result == "Время работы":
            table = working_hours
            self.ui.tableWidget.setColumnCount(6)
            collums = ['id', 'id_day', 'id_spec', 'work_hours', 'free_time', 'break_time']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                ["ID", "ID дня", "ID специализации", "рабочее время", "свободное время", "перерыв"])
        if result == "Приём":
            table = reception
            self.ui.tableWidget.setColumnCount(7)
            collums = ['id', 'id_patient', 'id_office', 'id_reception_status', 'id_personnel', 'id_spec', 'date']
            self.ui.tableWidget.setHorizontalHeaderLabels(
                ["ID", "ID пациента", "ID кабинета", "ID статуса приёма", "ID персонала", "id специализации", "Дата"])
        if result == "Мед книжка":
            table = med_knigа
            self.ui.tableWidget.setColumnCount(4)
            collums = ['id', 'id_patient', 'id_personnel', 'diagnoz']
            self.ui.tableWidget.setHorizontalHeaderLabels(["ID", "ID пациента", "ID персонала", "Диагноз"])
        line = session.query(table).count()
        self.ui.qTable = session.query(table).all()
        self.ui.tableWidget.setRowCount(line)
        sort = [0] * line
        j = 0
        p = session.query(specialization).all()
        for i in range(session.query(specialization).count()):
            self.ui.comboBox_5.addItem(str(p[i].id))
        for i in range(line):
            self.ui.comboBox_3.addItem(str(i + 1))
        for row, form in enumerate(self.ui.qTable):
            col = 0
            for c in collums:
                for k, v in vars(form).items():
                    if c == k:
                        self.ui.tableWidget.setItem(row, col, QTableWidgetItem(str(v)))
                        col += 1
        for i in session.query(table).all():
            sort[j] = i.id
            j += 1
        self.ui.comboBox_3.clear()
        sort.sort()
        for i in sort:
            self.ui.comboBox_3.addItem(str(i))
        perem = 0
        j = -1
        for i in sort:
            j += 1
            for row, form in enumerate(self.ui.qTable):
                col = 0
                for c in collums:
                    for k, v in vars(form).items():
                        if c == k:
                            if c == 'id':
                                perem = v
                            if ((c == 'id') & (v == i)) | ((c != 'id') & (perem == i)):
                                self.ui.tableWidget.setItem(j, col, QTableWidgetItem(str(v)))
                                col += 1
        self.ui.tableWidget.resizeColumnsToContents()

    def change_table(self):
        self.ui.statusbar.clearMessage()
        if self.ui.prov == 0:
            self.ui.prov = 1
        change = self.ui.comboBox_4.currentText()
        stlb = int(self.ui.comboBox_3.currentText())
        text = self.ui.textEdit.toPlainText()
        lis = text.split(',')
        if change == "Добавить строку":
            line = session.query(reception).count()
            tabl = session.query(reception).all()
            f = 0  # флаг
            free_id = 1  # свободный индекс
            while f == 0:
                for i in range(line):
                    if free_id == tabl[i].id:
                        f = 1
                if f == 1:
                    free_id = free_id + 1
                    f = 0
                else:
                    f = 1
            pers = session.query(personnel).filter(personnel.id == int(lis[0])).first()
            new = reception(id=free_id, id_patient=classes.glob_id, id_office=pers.id_office, id_reception_status=2, id_personnel=int(lis[0]), id_spec=pers.id_spec, date=lis[1])
            session.add(new)
        if change == "Отменить запись на приём":
            query = session.query(reception).get(stlb)
            if classes.glob_id == query.id_patient:
                query.id_reception_status = 3
            else:
                self.ui.statusbar.showMessage("Вы не можете отменить чужую запись на приём")
        self.write_table()

    def autoplay_spec(self):
        if self.ui.prov == 0:
            self.ui.prov = 1
        spec_id = int(self.ui.comboBox_5.currentText())
        dt = datetime.now()
        t2 = time(0)
        dt= datetime.combine(dt.date(), t2)
        while datetime.now() >= dt:
            dt = dt + timedelta(hours=1)
        tabl = session.query(reception).filter(reception.id_spec== spec_id).all()
        line = session.query(reception).filter(reception.id_spec== spec_id).count()
        f=0
        while f==0:
            f2=0
            while f2==0:
                for i in range(line):
                    if dt==datetime.strptime(tabl[i].date,'%Y-%m-%d %H:%M:%S'):
                        f2=1
                if f2==1:
                    f2=0
                    dt = dt + timedelta(hours=1)
                else:
                    f2=1
            wh = session.query(working_hours).filter(working_hours.id_spec == spec_id,working_hours.id_day==datetime.isoweekday(dt))
            line_wh = session.query(working_hours).filter(working_hours.id_spec == spec_id,working_hours.id_day==datetime.isoweekday(dt)).count()
            if line_wh!=0:
                for j in range(line_wh):
                    wh2=datetime.strptime(wh[j].work_hours[1],'%H:%M:%S')
                    wh1=datetime.strptime(wh[j].work_hours[0],'%H:%M:%S')
                    bt1 = datetime.strptime(wh[j].break_time[0], '%H:%M:%S')
                    bt2 = datetime.strptime(wh[j].break_time[1], '%H:%M:%S')
                    if (wh1.time()<=dt.time())&(wh2.time()>=dt.time())&((bt1.time()>dt.time())|(bt2.time()<dt.time())):
                        tabl2 =session.query(personnel).filter(personnel.id_spec == spec_id)
                        line2 = session.query(personnel).filter(personnel.id_spec == spec_id).count()
                        for k in range(line2):
                            line_l2 = session.query(reception).filter(reception.id_personnel == tabl2[k].id).count()
                            if line_l2 != 0:
                                l1 = session.query(reception).filter(reception.id_personnel == tabl2[k].id).first()
                                l2 = session.query(reception).filter(reception.id_office == l1.id_office,datetime.strptime(l1.date, '%Y-%m-%d %H:%M:%S')==dt).count()
                                if l2 == 0:
                                    res =dt
                                    d_per = tabl2[k].id
                                    n_c = tabl2[k].id_office
                                    f=1
                            else:
                                l1 = session.query(reception).filter(reception.id_office == tabl2[k].id_office).first()
                                l2 = session.query(reception).filter(reception.id_office == l1.id_office, datetime.strptime(l1.date, '%Y-%m-%d %H:%M:%S') == dt).count()
                                if l2 == 0:
                                    res =dt
                                    d_per = tabl2[k].id
                                    n_c = tabl2[k].id_office
                                    f=1
            dt = dt + timedelta(hours=1)
        tabl = session.query(reception).all()
        line = session.query(reception).count()
        f = 0  # флаг
        free_id = 1  # свободный индекс
        while f == 0:
            for i in range(line):
                if free_id == tabl[i].id:
                    f = 1
            if f == 1:
                free_id = free_id + 1
                f = 0
            else:
                f = 1
        new = reception(id=free_id, id_patient=classes.glob_id, id_office=n_c, id_reception_status=1, id_personnel=d_per, id_spec=spec_id, date=str(res))
        session.add(new)
        self.write_table()

    def confirm_change(self):
        if self.ui.prov != 0:
            session.commit()
            dump_sqlalchemy()
            self.ui.prov = 0
            self.write_table()
        else:
            self.ui.statusbar.showMessage("Сначала внесите изменение в таблицу")

    def undo_change(self):
        if self.ui.prov != 0:
            session.rollback()
            self.ui.prov = 0
            self.write_table()
        else:
            self.ui.statusbar.showMessage("Сначала внесите изменение в таблицу")



class log_panel(QMainWindow):
    def __init__(self, parent=None):
        super(log_panel, self).__init__(parent)
        self.ui = login_panell()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(lambda: self.inter_to_app())
        self.ui.pushButton_2.clicked.connect(lambda: self.create_acc())

    def inter_to_app(self):
        log = self.ui.textEdit.toPlainText()
        pasw = self.ui.textEdit_2.toPlainText()
        if log == "":
            self.ui.statusbar.showMessage("Введите логин")
            return
        if pasw == "":
            self.ui.statusbar.showMessage("Введите пароль")
            return
        for i in session.query(accounts).all():
            if log == i.login:
                if pasw == i.password:
                    self.hide()
                    if i.id_stat == 1:
                        dialog = admin_panel(parent=self)
                        dialog.show()
                    if i.id_stat == 2:
                        j = session.query(accounts).filter(accounts.login== log).first()
                        p = session.query(patient).filter(patient.id_acc== j.id).first()
                        classes.glob_id= p.id
                        dialog = patient_panel(parent=self)
                        dialog.show()
                    if i.id_stat == 3:
                        j = session.query(accounts).filter(accounts.login == log).first()
                        p = session.query(personnel).filter(personnel.id_acc == j.id).first()
                        classes.glob_id= p.id
                        dialog = personal_panel(parent=self)
                        dialog.show()
                else:
                    self.ui.statusbar.showMessage("Введён неверный пароль")
                    return
            self.ui.statusbar.showMessage("Введено несуществующее имя пользователя")

    def create_acc(self):
        log = self.ui.textEdit.toPlainText()
        pasw = self.ui.textEdit_2.toPlainText()
        if log == "":
            self.ui.statusbar.showMessage("Введите логин для нового пользователя")
            return
        if pasw == "":
            self.ui.statusbar.showMessage("Введите пароль для нового пользователя")
            return
        for i in session.query(accounts).all():
            if log == i.login:
                self.ui.statusbar.showMessage("Пользователь с таким логином уже существует")
                return
        n_id = int(self.ui.comboBox.currentText())
        line = session.query(accounts).count()
        tabl = session.query(accounts).all()
        f = 0  # флаг
        free_id = 1  # свободный индекс
        while f == 0:
            for i in range(line):
                if free_id == tabl[i].id:
                    f = 1
            if f == 1:
                free_id = free_id + 1
                f = 0
            else:
                f = 1
        if n_id == 1:
            user_setting = accounts(id=free_id, login=log, password=pasw, id_stat=1)
            session.add(user_setting)
            session.commit()
            dump_sqlalchemy()
            self.hide()
            dialog = admin_panel(parent=self)
            dialog.show()
        if n_id == 2:
            user_setting = accounts(id=free_id, login=log, password=pasw, id_stat=2)
            session.add(user_setting)
            session.commit()
            dump_sqlalchemy()
            f = 0  # флаг
            line = session.query(patient).count()
            tabl = session.query(patient).all()
            free_id2 = 1  # свободный индекс
            while f == 0:
                for i in range(line):
                    if free_id2 == tabl[i].id:
                        f = 1
                if f == 1:
                    free_id2 = free_id2 + 1
                    f = 0
                else:
                    f = 1
            user_setting = patient(id=free_id2, name=self.ui.textEdit_3.toPlainText(), id_acc=free_id)
            session.add(user_setting)
            session.commit()
            dump_sqlalchemy()

            self.hide()
            dialog = patient_panel(parent=self)
            dialog.show()

        if n_id == 3:
            user_setting = accounts(id=free_id, login=log, password=pasw, id_stat=3)
            session.add(user_setting)
            session.commit()
            dump_sqlalchemy()
            f = 0  # флаг
            line = session.query(personnel).count()
            tabl = session.query(personnel).all()
            free_id2 = 1  # свободный индекс
            while f == 0:
                for i in range(line):
                    if free_id2 == tabl[i].id:
                        f = 1
                if f == 1:
                    free_id2 = free_id2 + 1
                    f = 0
                else:
                    f = 1
            user_setting = personnel(id=free_id2, id_spec=int(self.ui.comboBox_2.currentText()), id_office= int(self.ui.comboBox_3.currentText()),name=self.ui.textEdit_3.toPlainText(), id_acc=free_id)
            session.add(user_setting)
            session.commit()
            dump_sqlalchemy()
            self.hide()
            dialog = personal_panel(parent=self)
            dialog.show()